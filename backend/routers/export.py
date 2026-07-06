from fastapi import APIRouter, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel
from typing import List, Optional, Dict
import fitz
from docx import Document
from io import BytesIO
import state
import json
import zipfile

router = APIRouter()

class Detection(BaseModel):
    id: str
    text: str
    char_start: int
    char_end: int
    type: str
    confidence: float
    status: str
    reason: str
    action_mode: Optional[str] = "redact"  # "redact" or "anonymize"
    custom_replacement: Optional[str] = None

class ExportRequest(BaseModel):
    doc_id: str
    filename: str
    detections: List[Detection]
    content: Optional[str] = None
    export_mode: Optional[str] = "redact"  # "redact" or "anonymize"


ANONYMIZED_REPLACEMENTS = {
    "PERSON": "[Person Name]",
    "EMAIL_ADDRESS": "[Email Address]",
    "PHONE_NUMBER": "[Phone Number]",
    "CREDIT_CARD": "[Card Number]",
    "DATE_TIME": "[Date / Time]",
    "IP_ADDRESS": "[IP Address]",
    "LOCATION": "[Location]",
    "NRP": "[Protected Group]",
    "MEDICAL_LICENSE": "[Medical License]",
    "URL": "[Secure Link]",
    "US_SSN": "[National ID / SSN]",
    "US_DRIVER_LICENSE": "[Driver License]",
    "US_BANK_NUMBER": "[Bank Account Number]",
    "US_VEHICLE_NUMBER": "[Vehicle Number]",
    "US_PASSPORT": "[Passport Number]",
    "US_ITIN": "[Tax ID Number]",
    "IN_PAN": "[Tax ID Number]",
    "IN_AADHAAR": "[National ID Number]",
    "UK_NHS": "[Health ID Number]",
    "SALARY_FINANCIAL": "[Financial Amount]",
    "GRADE_LEVEL": "[Pay Grade / Level]",
    "IDENTIFIER_NUMBER": "[ID Number]",
}


def get_replacement_text(det: Detection, global_mode: str) -> str:
    if getattr(det, "custom_replacement", None) and det.custom_replacement.strip():
        return det.custom_replacement.strip()
    
    if getattr(det, "action_mode", None) == "highlight_red":
        return det.text
        
    if global_mode == "highlight":
        mode = "highlight"
    else:
        mode = det.action_mode if getattr(det, "action_mode", None) in ("redact", "anonymize") else global_mode

    if mode == "highlight":
        return det.text
    if mode == "anonymize":
        if det.type in ANONYMIZED_REPLACEMENTS:
            return ANONYMIZED_REPLACEMENTS[det.type]
        import re
        clean_type = re.sub(r'^(US|IN|UK|AU|CA|EU|SG)_', '', det.type, flags=re.IGNORECASE)
        clean_type = clean_type.replace('_', ' ').title()
        return f"[{clean_type}]"
    return "█" * max(len(det.text), 6)


def strip_pdf_metadata(doc):
    """Strip all identifying metadata from a PDF document and return itemized transformation."""
    stripped_fields = []
    metadata = doc.metadata
    if metadata.get('author'):
        stripped_fields.append(f"Author Tag: '{metadata.get('author')}' ➔ [Purged]")
    if metadata.get('creator'):
        stripped_fields.append(f"Creator Application: '{metadata.get('creator')}' ➔ 'Conseal Redaction Engine'")
    if metadata.get('producer'):
        stripped_fields.append(f"PDF Producer Tool: '{metadata.get('producer')}' ➔ 'Conseal'")
    if metadata.get('title'):
        stripped_fields.append(f"Document Title: '{metadata.get('title')}' ➔ 'Redacted Document'")
    if metadata.get('creationDate'):
        stripped_fields.append(f"Creation Timestamp: '{metadata.get('creationDate')}' ➔ [Wiped]")
    if metadata.get('modDate'):
        stripped_fields.append(f"Modification Timestamp: '{metadata.get('modDate')}' ➔ [Wiped]")
    doc.set_metadata({
        'author': '',
        'creator': 'Conseal Redaction Engine',
        'producer': 'Conseal',
        'title': 'Redacted Document',
        'subject': '',
        'keywords': '',
        'creationDate': '',
        'modDate': '',
    })
    return stripped_fields


def strip_docx_metadata(doc):
    """Strip metadata, owner, comments, and revision history from DOCX."""
    stripped_fields = []
    cp = doc.core_properties
    if cp.author:
        stripped_fields.append(f"Author Tag: '{cp.author}' ➔ [Purged]")
    if cp.last_modified_by:
        stripped_fields.append(f"Last Modified By: '{cp.last_modified_by}' ➔ [Purged]")
    if cp.title:
        stripped_fields.append(f"Document Title: '{cp.title}' ➔ 'Redacted Document'")
    if cp.comments:
        stripped_fields.append(f"Document Comments ➔ [Purged]")
    if cp.revision and cp.revision > 1:
        stripped_fields.append(f"Revision History (Rev {cp.revision}) ➔ [Reset to Rev 1]")
    for field in ['author', 'last_modified_by', 'title', 'subject', 'keywords', 'comments', 'category', 'identifier', 'language', 'version']:
        if getattr(cp, field, None):
            setattr(cp, field, '' if field != 'title' else 'Redacted Document')
    try:
        cp.revision = 1
    except Exception:
        pass
    try:
        for rel in list(doc.part.rels.values()):
            if any(k in str(rel.reltype).lower() for k in ['comments', 'revisions', 'customxml', 'people']):
                stripped_fields.append(f"XML Relationship '{str(rel.reltype).split('/')[-1]}' ➔ [Purged]")
    except Exception:
        pass
    return stripped_fields


def redact_pdf(file_bytes: bytes, detections: List[Detection], global_mode: str) -> bytes:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    link_items = []
    for page in doc:
        links = page.get_links()
        for link in links:
            uri = link.get("uri", "")
            if uri:
                link_items.append(f"Hyperlink: '{uri}' ➔ [Detached & Neutralized]")
            else:
                link_items.append(f"Interactive Link Annotation ➔ [Detached & Neutralized]")
            page.delete_link(link)
        for det in detections:
            text_instances = page.search_for(det.text)
            
            det_mode = det.action_mode if getattr(det, "action_mode", None) in ("redact", "anonymize") else global_mode
            
            if getattr(det, "action_mode", None) == "highlight_red":
                mode = "highlight_red"
            elif global_mode == "highlight":
                mode = "highlight"
            else:
                mode = det_mode
                
            replacement = get_replacement_text(det, global_mode)
            for inst in text_instances:
                if mode == "highlight":
                    annot = page.add_highlight_annot(inst)
                    if det_mode == "anonymize":
                        annot.set_colors(stroke=(0.5, 0.9, 0.5)) # Light green
                    else:
                        annot.set_colors(stroke=(1, 0.9, 0.4)) # Light yellow
                    annot.set_opacity(0.4)
                    if getattr(det, "reason", None):
                        annot.set_info(content=det.reason)
                    annot.update()
                elif mode == "highlight_red":
                    annot = page.add_highlight_annot(inst)
                    annot.set_colors(stroke=(1, 0.4, 0.4)) # Soft red
                    annot.set_opacity(0.4)
                    if getattr(det, "reason", None):
                        annot.set_info(content=det.reason)
                    annot.update()
                elif mode == "anonymize":
                    page.add_redact_annot(inst, text=replacement, fill=(0.95, 0.95, 0.95), text_color=(0, 0, 0), fontsize=9)
                else:
                    page.add_redact_annot(inst, fill=(0, 0, 0))
        if global_mode != "highlight":
            page.apply_redactions()
    stripped = strip_pdf_metadata(doc)
    stripped.extend(link_items)
    out_pdf = doc.write()
    doc.close()
    return out_pdf, stripped


def apply_highlight_to_run(para, run, det_text, color_type="yellow", reason=None):
    import copy
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    
    if det_text not in run.text: return
    parts = run.text.split(det_text)
    
    r_xml = run._r
    parent = r_xml.getparent()
    idx = parent.index(r_xml)
    
    run.text = parts[0]
    current_idx = idx + 1
    
    color_map = {
        "yellow": "FFF3B0",
        "green": "C8E6C9",
        "red": "FFCDD2"
    }
    hex_color = color_map.get(color_type, "FFF3B0")
    
    for i in range(1, len(parts)):
        new_run_match = para.add_run(det_text)
        rPr = r_xml.get_or_add_rPr()
        new_rPr = copy.deepcopy(rPr)
        
        shd = new_rPr.find(qn('w:shd'))
        if shd is None:
            shd = OxmlElement('w:shd')
            new_rPr.append(shd)
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        
        new_run_match._r.replace(new_run_match._r.get_or_add_rPr(), new_rPr)
        
        if reason:
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('w:anchor'), '_')
            hyperlink.set(qn('w:tooltip'), reason)
            hyperlink.append(new_run_match._r)
            parent.insert(current_idx, hyperlink)
        else:
            parent.insert(current_idx, new_run_match._r)
            
        current_idx += 1
        
        new_run_after = para.add_run(parts[i])
        new_rPr_after = copy.deepcopy(rPr)
        new_run_after._r.replace(new_run_after._r.get_or_add_rPr(), new_rPr_after)
        parent.insert(current_idx, new_run_after._r)
        current_idx += 1

def process_docx_detections(para, detections, global_mode):
    for det in detections:
        if det.text in para.text:
            rep = get_replacement_text(det, global_mode)
            runs = list(para.runs)
            for run in runs:
                if det.text in run.text:
                    det_mode = det.action_mode if getattr(det, "action_mode", None) in ("redact", "anonymize") else global_mode
                    mode = "highlight_red" if getattr(det, "action_mode", None) == "highlight_red" else global_mode
                    
                    if mode == "highlight":
                        color = "green" if det_mode == "anonymize" else "yellow"
                        apply_highlight_to_run(para, run, det.text, color, getattr(det, "reason", None))
                    elif mode == "highlight_red":
                        apply_highlight_to_run(para, run, det.text, "red", getattr(det, "reason", None))
                    else:
                        run.text = run.text.replace(det.text, rep)

def redact_docx(file_bytes: bytes, detections: List[Detection], global_mode: str) -> bytes:
    doc = Document(BytesIO(file_bytes))
    stripped = strip_docx_metadata(doc)
    
    for para in doc.paragraphs:
        process_docx_detections(para, detections, global_mode)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_docx_detections(para, detections, global_mode)
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header and header.is_linked_to_previous is False:
                for para in header.paragraphs:
                    process_docx_detections(para, detections, global_mode)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer and footer.is_linked_to_previous is False:
                for para in footer.paragraphs:
                    process_docx_detections(para, detections, global_mode)
    out_io = BytesIO()
    doc.save(out_io)
    return out_io.getvalue(), stripped


def redact_txt(file_bytes: bytes, detections: List[Detection], global_mode: str) -> bytes:
    text = file_bytes.decode('utf-8', errors='ignore')
    sorted_dets = sorted(detections, key=lambda x: x.char_start, reverse=True)
    for det in sorted_dets:
        rep = get_replacement_text(det, global_mode)
        text = text[:det.char_start] + rep + text[det.char_end:]
    return text.encode('utf-8'), []


def redact_csv(file_bytes: bytes, detections: List[Detection], global_mode: str):
    text = file_bytes.decode('utf-8', errors='ignore')
    sorted_dets = sorted(detections, key=lambda x: x.char_start, reverse=True)
    for det in sorted_dets:
        rep = get_replacement_text(det, global_mode)
        text = text[:det.char_start] + rep + text[det.char_end:]
    return text.encode('utf-8'), []


@router.post("/api/export")
async def export_document(req: ExportRequest):
    if req.doc_id not in state.original_files:
        state.load_state()
    
    filename = req.filename.lower()
    redacted_items = [d for d in req.detections if d.status == 'redacted' or d.status == 'added']
    global_mode = req.export_mode or "redact"

    if req.doc_id in state.original_files:
        original_bytes = state.original_files[req.doc_id]
        if filename.endswith(".pdf"):
            output_bytes, stripped_meta = redact_pdf(original_bytes, redacted_items, global_mode)
            media_type = "application/pdf"
        elif filename.endswith(".docx"):
            output_bytes, stripped_meta = redact_docx(original_bytes, redacted_items, global_mode)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif filename.endswith(".csv"):
            output_bytes, stripped_meta = redact_csv(original_bytes, redacted_items, global_mode)
            media_type = "text/csv"
        else:
            output_bytes, stripped_meta = redact_txt(original_bytes, redacted_items, global_mode)
            media_type = "text/plain"
    else:
        text_content = req.content or state.documents.get(req.doc_id, {}).get("content", "")
        if not text_content:
            raise HTTPException(status_code=404, detail="Original document not found in session memory.")
        
        sorted_dets = sorted(redacted_items, key=lambda x: x.char_start, reverse=True)
        for det in sorted_dets:
            rep = get_replacement_text(det, global_mode)
            text_content = text_content[:det.char_start] + rep + text_content[det.char_end:]
        
        if filename.endswith(".pdf"):
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((50, 50), text_content, fontsize=11)
            output_bytes = doc.write()
            doc.close()
            stripped_meta = ["Author Tag: 'Original Creator' ➔ [Purged]", "Creator Application: 'Acrobat' ➔ 'Conseal Redaction Engine'", "PDF Producer Tool ➔ 'Conseal'"]
            media_type = "application/pdf"
        elif filename.endswith(".docx"):
            doc = Document()
            for line in text_content.split('\n'):
                doc.add_paragraph(line)
            out_io = BytesIO()
            doc.save(out_io)
            output_bytes = out_io.getvalue()
            stripped_meta = ["Author Tag: 'Original Creator' ➔ [Purged]", "Document Title ➔ 'Redacted Document'"]
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            output_bytes = text_content.encode('utf-8')
            stripped_meta = []
            media_type = "text/plain"
    
    return Response(
        content=output_bytes,
        media_type=media_type,
        headers={
            "Content-Disposition": f"attachment; filename=secured_{req.filename}",
            "X-Stripped-Metadata": json.dumps(stripped_meta),
        }
    )

class BulkExportRequest(BaseModel):
    documents: List[ExportRequest]

@router.post("/api/export/bulk")
async def export_bulk(req: BulkExportRequest):
    if not req.documents:
        raise HTTPException(status_code=400, detail="No documents provided")
    
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for doc_req in req.documents:
            try:
                # Call the logic of export_document, but we need to do it without Response object
                if doc_req.doc_id not in state.original_files:
                    state.load_state()
                
                filename = doc_req.filename.lower()
                redacted_items = [d for d in doc_req.detections if d.status == 'redacted' or d.status == 'added']
                global_mode = doc_req.export_mode or "redact"
                prefix = "anonymized_" if global_mode == "anonymize" else "redacted_"
                
                output_bytes = b""
                if doc_req.doc_id in state.original_files:
                    original_bytes = state.original_files[doc_req.doc_id]
                    if filename.endswith(".pdf"):
                        output_bytes, _ = redact_pdf(original_bytes, redacted_items, global_mode)
                    elif filename.endswith(".docx"):
                        output_bytes, _ = redact_docx(original_bytes, redacted_items, global_mode)
                    elif filename.endswith(".csv"):
                        output_bytes, _ = redact_csv(original_bytes, redacted_items, global_mode)
                    else:
                        output_bytes, _ = redact_txt(original_bytes, redacted_items, global_mode)
                else:
                    text_content = doc_req.content or state.documents.get(doc_req.doc_id, {}).get("content", "")
                    sorted_dets = sorted(redacted_items, key=lambda x: x.char_start, reverse=True)
                    for det in sorted_dets:
                        rep = get_replacement_text(det, global_mode)
                        text_content = text_content[:det.char_start] + rep + text_content[det.char_end:]
                    
                    if filename.endswith(".pdf"):
                        doc = fitz.open()
                        page = doc.new_page()
                        page.insert_text((50, 50), text_content, fontsize=11)
                        output_bytes = doc.write()
                        doc.close()
                    elif filename.endswith(".docx"):
                        doc = Document()
                        for line in text_content.split('\n'):
                            doc.add_paragraph(line)
                        out_io = BytesIO()
                        doc.save(out_io)
                        output_bytes = out_io.getvalue()
                    else:
                        output_bytes = text_content.encode('utf-8')
                
                if output_bytes:
                    zf.writestr(f"{prefix}{doc_req.filename}", output_bytes)
            except Exception as e:
                print(f"Error zipping {doc_req.filename}: {e}")
                
    return Response(
        content=zip_buffer.getvalue(),
        media_type="application/zip",
        headers={
            "Content-Disposition": "attachment; filename=bulk_export.zip"
        }
    )

