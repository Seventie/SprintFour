from fastapi import APIRouter, UploadFile, File, Form
import fitz  # PyMuPDF
from docx import Document
import uuid
from typing import List, Optional
from io import BytesIO
import state
import csv
from services.heuristic import run_heuristic_detection, merge_detections, align_detection_boundaries
from services.gliner_service import detect_pii_gliner

router = APIRouter()


# --- Confidence-based status assignment ---
def assign_status(confidence: float) -> str:
    if confidence >= 0.85:
        return "redacted"
    elif confidence >= 0.50:
        return "missed"
    else:
        return "fp"


# --- Contextual reason generation ---
REASON_TEMPLATES = {
    "PERSON": "Presidio NLP identified a human personal name",
    "EMAIL_ADDRESS": "Regex/NLP matched standard RFC email address format",
    "PHONE_NUMBER": "Pattern matched standard telephone number structure",
    "CREDIT_CARD": "Luhn algorithm validated payment card number format",
    "DATE_TIME": "Entity recognizer flagged temporal/date reference",
    "IP_ADDRESS": "Regex matched IPv4/IPv6 network address format",
    "LOCATION": "Named Entity Recognition flagged geographic place/address",
    "NRP": "Pattern matched nationality, religious, or political group reference",
    "MEDICAL_LICENSE": "Pattern matched medical license number format",
    "URL": "Regex matched web URL/URI protocol format",
    "US_SSN": "Pattern matched United States Social Security Number",
    "US_DRIVER_LICENSE": "Pattern matched state driver license ID number",
}


def get_reason(entity_type: str, confidence: float, text: str) -> str:
    base = REASON_TEMPLATES.get(entity_type, f"Presidio NLP model detected {entity_type}")
    if confidence >= 0.85:
        return f"{base}. High confidence — auto-redacted for safety."
    elif confidence >= 0.50:
        return f"{base}. Medium confidence — flagged for human review."
    else:
        return f"{base}. Low confidence — likely a false positive."


def extract_text_from_pdf(file_bytes):
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text("text") + "\n"
    return text


def extract_text_from_docx(file_bytes):
    doc = Document(BytesIO(file_bytes))
    text = "\n".join([para.text for para in doc.paragraphs])
    return text


def extract_text_from_csv(file_bytes):
    text = file_bytes.decode('utf-8', errors='ignore')
    lines = []
    reader = csv.reader(text.splitlines())
    for row in reader:
        lines.append(" | ".join([cell.strip() for cell in row if cell.strip()]))
    return "\n".join(lines)


def extract_text_from_txt(file_bytes):
    return file_bytes.decode('utf-8', errors='ignore')


@router.post("/api/upload")
async def upload_files(
    files: List[UploadFile] = File(...),
    file_modes: Optional[str] = Form(None),
    policy: Optional[str] = Form("mixed_general")
):
    print(f"--- UPLOAD CALLED --- ({len(files)} files, policy={policy})")
    modes_dict = {}
    if file_modes:
        try:
            import json
            modes_dict = json.loads(file_modes)
        except Exception as e:
            print("Failed parsing file_modes:", e)

    documents = []
    all_detections = {}

    for idx, file in enumerate(files):
        contents = await file.read()
        print(f"Processing: {file.filename} ({len(contents)} bytes)")
        filename = file.filename.lower()

        text = ""
        file_type = "txt"

        metadata = {}

        if filename.endswith(".pdf"):
            text = extract_text_from_pdf(contents)
            file_type = "pdf"
            try:
                pdf_doc = fitz.open(stream=contents, filetype="pdf")
                md = pdf_doc.metadata or {}
                metadata = {
                    "author": md.get("author", ""),
                    "creator": md.get("creator", ""),
                    "title": md.get("title", ""),
                    "created": md.get("creationDate", ""),
                    "modified": md.get("modDate", ""),
                    "tool": md.get("producer", ""),
                }
                pdf_doc.close()
            except Exception as e:
                print(f"PDF metadata extraction error: {e}")
        elif filename.endswith(".docx"):
            text = extract_text_from_docx(contents)
            file_type = "docx"
            try:
                docx_doc = Document(BytesIO(contents))
                cp = docx_doc.core_properties
                metadata = {
                    "author": cp.author or "",
                    "creator": cp.last_modified_by or "",
                    "title": cp.title or "",
                    "created": str(cp.created) if cp.created else "",
                    "modified": str(cp.modified) if cp.modified else "",
                    "tool": "",
                }
            except Exception as e:
                print(f"DOCX metadata extraction error: {e}")
        elif filename.endswith(".csv"):
            text = extract_text_from_csv(contents)
            file_type = "csv"
        else:
            text = extract_text_from_txt(contents)
            file_type = "txt"

        doc_id = str(uuid.uuid4())
        state.original_files[doc_id] = contents

        mode = modes_dict.get(str(idx)) or modes_dict.get(file.filename) or "redact"

        # --- Core NLP Detection (GLiNER2-PII) totally replacing Presidio ---
        model_dets = detect_pii_gliner(text, policy_name=policy or "mixed_general")

        heuristic_dets = run_heuristic_detection(text)
        merged_dets = merge_detections(model_dets, heuristic_dets)
        merged_dets = align_detection_boundaries(text, merged_dets)

        for det in merged_dets:
            det["action_mode"] = mode
            src = det.get("source", "unknown")
            print(f"    [{src}] {det['type']} | '{det['text'][:30]}' | {det['confidence']} | {det['status']} | {mode}")

        doc_record = {
            "doc_id": doc_id,
            "filename": file.filename,
            "file_type": file_type,
            "status": "ready",
            "char_count": len(text),
            "content": text,
            "default_action_mode": mode,
            "metadata": metadata,
        }

        state.documents[doc_id] = doc_record
        state.detections[doc_id] = merged_dets

        documents.append(doc_record)
        all_detections[doc_id] = merged_dets

    state.save_state()

    return {
        "documents": documents,
        "detections": all_detections,
    }

